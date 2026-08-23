"""Render a markdown résumé as a .docx.

The generated and per-job tailored résumés were produced as markdown and
uploaded as ``.md``. That file is what the apply engine then attaches to a
Greenhouse / Lever / Ashby résumé field — and those fields accept pdf, doc,
docx, txt and rtf. A ``.md`` upload is rejected by the form, so the application
bounced and the engine parked it as unconfirmed. Every "build one for me" user
and everyone on the tailored strategy hit it.

.docx rather than PDF because it needs no new dependency (python-docx is
already here for reading uploads) and because ATS résumé parsers handle Word
better than a generated PDF's text layer.

The markdown these produce is our own (``resume_builder.build_markdown``) or an
LLM's, so this handles the small subset both emit — headings, bullets, bold —
and treats anything else as a paragraph rather than failing.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Inline spans, matched in one pass so a run is never emitted twice.
# Order matters: ** before * so bold is not read as two italics.
_INLINE_RE = re.compile(
    r"\*\*(?P<bold>.+?)\*\*"
    r"|__(?P<bold2>.+?)__"
    r"|\*(?P<em>[^*\n]+?)\*"
    r"|(?<![A-Za-z0-9])_(?P<em2>[^_\n]+?)_(?![A-Za-z0-9])"
    r"|`(?P<code>[^`\n]+?)`"
    r"|\[(?P<text>[^\]\n]+?)\]\((?P<href>[^)\n]+?)\)"
)
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(?P<text>\S.*)$")
_NUMBER_RE = re.compile(r"^\s*\d{1,3}[.)]\s+(?P<text>\S.*)$")
_HEADING_RE = re.compile(r"^(?P<hashes>#{1,6})\s+(?P<text>\S.*)$")
_QUOTE_RE = re.compile(r"^\s*>\s*(?P<text>.*)$")
# ---, ***, ___ and longer. A horizontal rule is layout, not content.
_RULE_RE = re.compile(r"^\s*([-*_])\1{2,}\s*$")


def _add_runs(paragraph, text: str) -> None:
    """Write text into a paragraph, rendering inline markdown as formatting.

    Anything left as literal syntax here ends up in a document a recruiter
    reads: the tailored résumé is written by a model, and models emit
    ``*emphasis*``, ``` `code` ``` and ``[links](urls)`` freely.
    """
    position = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        groups = match.groupdict()
        if groups["bold"] or groups["bold2"]:
            paragraph.add_run(groups["bold"] or groups["bold2"]).bold = True
        elif groups["em"] or groups["em2"]:
            paragraph.add_run(groups["em"] or groups["em2"]).italic = True
        elif groups["code"]:
            # Not a monospace run: an ATS parser reads the text, and a résumé
            # has no reason to look like a terminal.
            paragraph.add_run(groups["code"])
        elif groups["text"]:
            # Keep the URL — a portfolio link is the point of the link.
            label, href = groups["text"], groups["href"]
            paragraph.add_run(label if label == href else f"{label} ({href})")
        position = match.end()
    remaining = text[position:]
    if remaining or not paragraph.runs:
        paragraph.add_run(remaining)


def markdown_to_docx(markdown: str) -> bytes:
    """Convert résumé markdown to .docx bytes."""
    document = Document()
    # A résumé is read by a parser first and a person second: one common family,
    # no columns, no tables, nothing that confuses text extraction.
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for raw_line in (markdown or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        # A horizontal rule is layout the docx expresses with spacing, and
        # "---" printed in the middle of a résumé just looks like a mistake.
        if _RULE_RE.match(line):
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            level = len(heading.group("hashes"))
            text = heading.group("text").strip()
            if level == 1:
                # The name. Centred, and the only thing that gets to be large.
                paragraph = document.add_heading(level=0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_runs(paragraph, text)
            else:
                _add_runs(document.add_heading(level=min(level, 4)), text)
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            _add_runs(document.add_paragraph(style="List Bullet"), bullet.group("text"))
            continue

        numbered = _NUMBER_RE.match(line)
        if numbered:
            _add_runs(document.add_paragraph(style="List Number"), numbered.group("text"))
            continue

        quote = _QUOTE_RE.match(line)
        if quote:
            # Strip the marker; a stray ">" reads as a typo on a résumé.
            text = quote.group("text").strip()
            if not text:
                continue
            _add_runs(document.add_paragraph(), text)
            continue

        # A bare "#" or "-" with nothing after it is a stray marker, not content.
        if line.strip().strip("#-*+>") == "":
            continue

        _add_runs(document.add_paragraph(), line.strip())

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
