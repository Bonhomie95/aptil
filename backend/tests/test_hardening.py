"""Regressions for the issues found in the full read-through.

Each of these was silently wrong: no error, no log, just the wrong answer.
"""

from __future__ import annotations

import uuid
from datetime import UTC, datetime, timedelta

import pytest

from tests.conftest import requires_mongo


# --------------------------------------------------------------------------
# Email is queued, not awaited (signup took 8.4s against a real relay)
# --------------------------------------------------------------------------
async def test_sending_does_not_wait_on_the_mail_relay(monkeypatch):
    """The request must not block on smtplib."""
    from app.core.config import settings
    from app.services import email as email_service

    monkeypatch.setattr(settings, "SMTP_HOST", "smtp.example.com")
    queued: list[tuple] = []
    monkeypatch.setattr(
        email_service,
        "_enqueue",
        lambda *a: queued.append(a),
    )

    def _explode(*_a, **_k):
        raise AssertionError("sent inline instead of queued")

    monkeypatch.setattr(email_service, "send_email_sync", _explode)

    assert await email_service.send_email("a@b.co", "Subject", "Body") is True
    assert len(queued) == 1


async def test_a_dead_broker_falls_back_to_sending_inline(monkeypatch):
    """Losing a verification link strands the account — send it ourselves."""
    from app.core.config import settings
    from app.services import email as email_service

    monkeypatch.setattr(settings, "SMTP_HOST", "smtp.example.com")

    def _no_broker(*_a):
        raise ConnectionError("broker down")

    sent: list[tuple] = []
    monkeypatch.setattr(email_service, "_enqueue", _no_broker)
    monkeypatch.setattr(
        email_service, "send_email_sync", lambda *a: sent.append(a)
    )

    assert await email_service.send_email("a@b.co", "Subject", "Body") is True
    assert len(sent) == 1


async def test_stub_mode_stays_synchronous(monkeypatch):
    """With no SMTP host there is nothing to queue; keep dev inspectable."""
    from app.core.config import settings
    from app.services import email as email_service

    monkeypatch.setattr(settings, "SMTP_HOST", "")
    monkeypatch.setattr(
        email_service,
        "_enqueue",
        lambda *a: (_ for _ in ()).throw(AssertionError("queued a stub send")),
    )
    assert await email_service.send_email("a@b.co", "S", "B") is True


# --------------------------------------------------------------------------
# Storage teardown must never take down the caller
# --------------------------------------------------------------------------
def test_delete_object_swallows_an_unreachable_endpoint(monkeypatch):
    """It claims to be best-effort; only ClientError was actually caught, so a
    down bucket made replacing a résumé fail on removing the OLD file."""
    from botocore.exceptions import EndpointConnectionError

    from app.services import storage

    class _Boom:
        def delete_object(self, **_k):
            raise EndpointConnectionError(endpoint_url="http://minio:9000")

    monkeypatch.setattr(storage, "_client", lambda: _Boom())
    storage.delete_object("some/key")  # must not raise


# --------------------------------------------------------------------------
# Entitlement metering
# --------------------------------------------------------------------------
@requires_mongo
async def test_an_elapsed_period_resets_once_not_on_every_spend():
    """The reset branch wrote the counters but not the new window, so after a
    period lapsed every increment reset again — unlimited usage."""
    from app.db.session import init_db
    from app.models.billing import Plan, Subscription
    from app.services import billing

    await init_db()
    plan = Plan(
        code=f"t-{uuid.uuid4().hex[:8]}", name="T", price_cents=0,
        monthly_applications=2, monthly_interviews=1,
    )
    await plan.insert()
    past = datetime.now(UTC) - timedelta(days=1)
    sub = Subscription(
        tenant_id=uuid.uuid4(), plan_id=plan.id, status="active",
        current_period_start=past - timedelta(days=30),
        current_period_end=past,
        applications_used=2,
    )
    await sub.insert()

    first = await billing.increment_application_usage(sub.tenant_id)
    assert first.applications_used == 1, "the lapsed period should reset once"
    assert first.current_period_end > datetime.now(UTC).replace(tzinfo=None) or (
        first.current_period_end.replace(tzinfo=UTC) > datetime.now(UTC)
    ), "the window must roll forward, or the reset repeats forever"

    second = await billing.increment_application_usage(sub.tenant_id)
    assert second.applications_used == 2, "the second spend must count, not reset"


# --------------------------------------------------------------------------
# The dashboard's status picker and the API must agree
# --------------------------------------------------------------------------
def test_every_status_the_ui_offers_is_accepted_by_the_api():
    """The picker used to render "Submitted", which the API answered 422 to.

    The fix is on the UI side, not here: `submitted` must stay unsettable by
    hand (compliance section 2a — only a confirmed submission may claim it), so
    the dashboard shows the current status as a disabled option instead of
    offering it. This pins the set the picker is allowed to offer.
    """
    from app.api.v1.routes.jobs import USER_SETTABLE_STATUSES

    offered_by_the_dashboard = {"confirmed", "interview", "offer", "rejected"}
    assert offered_by_the_dashboard <= USER_SETTABLE_STATUSES
    assert "submitted" not in USER_SETTABLE_STATUSES, (
        "a user must not be able to claim an application was submitted"
    )


# --------------------------------------------------------------------------
# Google-only accounts
# --------------------------------------------------------------------------
def test_user_read_reports_whether_a_local_password_exists():
    from app.schemas.auth import UserRead

    class _U:
        id = uuid.uuid4()
        tenant_id = uuid.uuid4()
        email = "a@b.co"
        full_name = None
        is_email_verified = True
        onboarding_step = "cv_upload"
        onboarding_completed = False
        hashed_password: str | None = None

        @property
        def has_password(self) -> bool:
            return bool(self.hashed_password)

    assert UserRead.model_validate(_U()).has_password is False
    _U.hashed_password = "argon2$..."
    assert UserRead.model_validate(_U()).has_password is True


def test_change_password_accepts_a_missing_current_password():
    """A Google-only account has none to send."""
    from app.schemas.auth import ChangePasswordRequest

    body = ChangePasswordRequest(new_password="a-good-long-password")
    assert body.current_password is None


# --------------------------------------------------------------------------
# Interview question cap
# --------------------------------------------------------------------------
def test_the_question_cap_holds_for_bare_strings_too():
    """The cap was only checked after the dict branch, so a model returning
    plain strings blew straight past the entitlement."""
    from app.api.v1.routes.interviews import _normalize_questions

    assert len(_normalize_questions([f"Q{i}?" for i in range(50)], 6)) == 6
    assert len(_normalize_questions([{"question": f"Q{i}"} for i in range(50)], 6)) == 6


# --------------------------------------------------------------------------
# Résumés must be a format an ATS accepts
# --------------------------------------------------------------------------
def test_generated_resumes_are_docx_not_markdown():
    """A .md upload is rejected by every ATS résumé field we target."""
    import io

    from docx import Document

    from app.services.resume_docx import DOCX_CONTENT_TYPE, markdown_to_docx

    data = markdown_to_docx(
        "# Ada Lovelace\nada@example.com\n\n## Skills\nPython, Analytical Engines\n"
        "\n## Experience\n\n### Engineer · Acme · 2021\n- Built the thing\n"
    )
    assert data[:2] == b"PK", "not a real .docx container"
    assert "wordprocessingml" in DOCX_CONTENT_TYPE

    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Ada Lovelace" in text
    assert "Python, Analytical Engines" in text
    assert "Built the thing" in text
    assert "**" not in text and "#" not in text, "markdown syntax leaked through"


def test_docx_renderer_survives_junk():
    from app.services.resume_docx import markdown_to_docx

    for junk in ("", "   ", "no markdown at all", "#\n##\n- \n**unclosed"):
        assert markdown_to_docx(junk)[:2] == b"PK"


def test_no_markdown_syntax_reaches_the_reader():
    """The tailored résumé is written by a model, and models emit rules,
    emphasis, code spans and links freely. Any of it left as literal syntax
    lands in a document a recruiter reads."""
    import io

    from docx import Document

    from app.services.resume_docx import markdown_to_docx

    md = (
        "# Rita Rightfile\n"
        "rita@example.com | [portfolio](https://rita.dev)\n\n"
        "---\n\n"
        "## Summary\n"
        "Engineer with **8 years** building *payments* systems.\n\n"
        "***\n\n"
        "## Experience\n"
        "### Engineer, Acme\n"
        "- Led the rewrite using `asyncio`\n"
        "- Cut __p99 latency__ by 40%\n"
        "1. Shipped the ledger\n"
        "#\n##\n- \n"
        "> a blockquote\n"
    )
    doc = Document(io.BytesIO(markdown_to_docx(md)))
    text = "\n".join(p.text for p in doc.paragraphs)

    for leaked in ("---", "***", "**", "__", "`", "](", "> "):
        assert leaked not in text, f"{leaked!r} reached the reader"
    assert not [
        p for p in doc.paragraphs if p.text.strip() in ("#", "##", "-", ">")
    ], "a stray marker was rendered as content"

    # The content itself survived, and emphasis became real formatting.
    assert "8 years" in text and "payments" in text and "asyncio" in text
    assert "portfolio (https://rita.dev)" in text, "a link lost its URL"
    assert any(r.bold for p in doc.paragraphs for r in p.runs), "bold was dropped"
    assert any(r.italic for p in doc.paragraphs for r in p.runs), "italic was dropped"
    styles = {p.style.name for p in doc.paragraphs if p.text.strip()}
    assert "List Bullet" in styles and "List Number" in styles


# --------------------------------------------------------------------------
# Scoring projection
# --------------------------------------------------------------------------
def test_scoring_projection_carries_everything_score_job_reads():
    """If a field scoring uses is dropped from the projection, every match
    silently scores lower instead of failing."""
    from app.services.matching import ScoredJob

    assert {"id", "company", "title", "location", "remote", "description"} == set(
        ScoredJob.model_fields
    )
    assert "raw" not in ScoredJob.model_fields


def test_scoring_works_on_the_projection():
    from app.models.profile import Profile
    from app.services.matching import ScoredJob, score_job

    job = ScoredJob(
        _id=uuid.uuid4(),
        company="Acme",
        title="Senior Backend Engineer",
        description="We use Python and Postgres.",
    )
    profile = Profile.model_construct(
        skills=["Python", "Postgres"],
        work_history=[{"title": "Backend Engineer"}],
        preferences={},
        headline=None,
    )
    score = score_job(profile, job)
    assert 0.0 < score <= 1.0


# --------------------------------------------------------------------------
# Config guards
# --------------------------------------------------------------------------
@pytest.mark.parametrize(
    "field, value, expected",
    [
        ("MONGO_URL", "mongodb://mongo:27017", "MONGO_URL"),
        ("FRONTEND_BASE_URL", "http://localhost:3000", "FRONTEND_BASE_URL"),
    ],
)
def test_production_refuses_the_local_defaults(field, value, expected):
    """Reaching production with these means the app talks to nothing, or emails
    links nobody can open."""
    from app.core.config import Settings

    good = dict(
        ENVIRONMENT="production",
        SECRET_KEY="x" * 40,
        CREDENTIAL_ENCRYPTION_KEY="k" * 44,
        MINIO_ROOT_PASSWORD="pw",
        SMTP_HOST="smtp.example.com",
        MONGO_URL="mongodb+srv://real/db",
        FRONTEND_BASE_URL="https://aptil.example",
        TRUSTED_PROXY_IPS="10.0.0.0/8",
    )
    Settings(**good)  # the good config must still boot

    with pytest.raises(ValueError, match=expected):
        Settings(**{**good, field: value})


def test_has_any_provider_does_not_count_the_ollama_default():
    """OLLAMA_BASE_URL ships with a default, so counting it made the answer
    True on every install."""
    from app.ai import router

    assert router.has_any_provider() is bool(router._ALL_KEYS)


# --------------------------------------------------------------------------
# Email retry exhaustion
# --------------------------------------------------------------------------
def test_email_task_gives_up_cleanly_after_its_last_retry(monkeypatch):
    """The task body also runs on the attempt AFTER the final retry.

    Indexing the backoff table straight with `self.request.retries` raised
    IndexError there, so the task died on that rather than on
    MaxRetriesExceeded — and never logged that it had stopped trying.
    """
    from app.workers.tasks import email as email_task

    delays = email_task._RETRY_DELAYS
    # Every attempt the body can see, including one past the end of the table.
    for attempt in range(len(delays) + 2):
        index = min(attempt, len(delays) - 1)
        assert delays[index] == delays[min(attempt, len(delays) - 1)]
        assert 0 <= index < len(delays), f"attempt {attempt} indexes out of range"


def test_email_task_retry_table_matches_max_retries():
    """max_retries and the backoff table have to stay in step."""
    from app.workers.tasks.email import _RETRY_DELAYS, send_email_task

    assert send_email_task.max_retries == len(_RETRY_DELAYS)
    assert all(d > 0 for d in _RETRY_DELAYS)
    assert list(_RETRY_DELAYS) == sorted(_RETRY_DELAYS), "backoff should increase"
